from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches

from .paths import ROOT, archive_dir, load_config, resolve, template_path
from .template import (
    REQUIRED_PLACEHOLDERS,
    build_master_template,
    docx_to_dotx,
    validate_template,
)


LOGO_NAMES = ("logo.png", "logo.jpg", "logo.jpeg", "logo.webp", "logo.svg")
ALLOWED_TEMPLATE_SUFFIXES = {".docx", ".dotx"}
ALLOWED_LOGO_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif"}


def images_dir() -> Path:
    return resolve(load_config()["paths"]["images"])


def logo_path() -> Path | None:
    folder = images_dir()
    for name in LOGO_NAMES:
        candidate = folder / name
        if candidate.exists():
            return candidate
    # Any image named logo.* 
    matches = sorted(folder.glob("logo.*"))
    return matches[0] if matches else None


def template_docx_path() -> Path:
    return resolve(load_config()["paths"]["template_fallback"])


def template_dotx_path() -> Path:
    return resolve(load_config()["paths"]["template"])


def _archive_file(path: Path) -> Path | None:
    if not path.exists():
        return None
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    dest = archive_dir() / f"{path.stem}-{stamp}{path.suffix}"
    archive_dir().mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, dest)
    return dest


def template_status() -> dict[str, Any]:
    logo = logo_path()
    active = None
    active_error = None
    try:
        active = template_path()
    except FileNotFoundError as exc:
        active_error = str(exc)

    validation = None
    if active is not None:
        try:
            validation = validate_template(active)
            validation = {
                "ok": validation["ok"],
                "missing": validation["missing"],
                "is_dotx": validation["is_dotx"],
                "path": str(Path(validation["path"]).relative_to(ROOT)),
            }
        except Exception as exc:  # noqa: BLE001
            validation = {"ok": False, "missing": [], "error": str(exc)}

    return {
        "active": str(active.relative_to(ROOT)) if active else None,
        "active_error": active_error,
        "docx": str(template_docx_path().relative_to(ROOT)),
        "docx_exists": template_docx_path().exists(),
        "dotx": str(template_dotx_path().relative_to(ROOT)),
        "dotx_exists": template_dotx_path().exists(),
        "logo": str(logo.relative_to(ROOT)) if logo else None,
        "logo_exists": logo is not None,
        "logo_url": "/api/template/logo" if logo else None,
        "download_docx": "/api/template/download?format=docx",
        "download_dotx": "/api/template/download?format=dotx",
        "required_placeholders": list(REQUIRED_PLACEHOLDERS),
        "validation": validation,
    }


def save_uploaded_template(data: bytes, filename: str) -> dict[str, Any]:
    """
    Accept a .docx or .dotx upload from the web UI (no repo access required).

    Saves both the working .docx twin and the gold-master .dotx.
    """
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_TEMPLATE_SUFFIXES:
        raise ValueError("Upload a .docx or .dotx file")
    if not data:
        raise ValueError("Empty upload")

    Template = resolve("Template")
    Template.mkdir(parents=True, exist_ok=True)
    docx_dest = template_docx_path()
    dotx_dest = template_dotx_path()

    _archive_file(docx_dest)
    _archive_file(dotx_dest)

    staging = Template / f"_upload_staging{suffix}"
    staging.write_bytes(data)

    try:
        if suffix == ".docx":
            shutil.copy2(staging, docx_dest)
            docx_to_dotx(docx_dest, dotx_dest)
        else:
            # Uploaded .dotx — keep as gold master; also write a .docx twin
            shutil.copy2(staging, dotx_dest)
            # Convert content-type back for editable twin by flipping types
            from .template import DOCX_MAIN_CONTENT_TYPE, DOTX_MAIN_CONTENT_TYPE
            import io
            import zipfile

            buffer = io.BytesIO()
            with zipfile.ZipFile(staging, "r") as zin, zipfile.ZipFile(
                buffer, "w", compression=zipfile.ZIP_DEFLATED
            ) as zout:
                for info in zin.infolist():
                    chunk = zin.read(info.filename)
                    if info.filename == "[Content_Types].xml":
                        text = chunk.decode("utf-8").replace(
                            DOTX_MAIN_CONTENT_TYPE, DOCX_MAIN_CONTENT_TYPE
                        )
                        chunk = text.encode("utf-8")
                    zout.writestr(info, chunk)
            docx_dest.write_bytes(buffer.getvalue())

        # Validate placeholders after install
        result = validate_template(dotx_dest if dotx_dest.exists() else docx_dest)
        if not result["ok"]:
            raise ValueError(
                "Template saved but missing required placeholders: "
                + ", ".join(result["missing"])
            )
    finally:
        if staging.exists():
            staging.unlink()

    return template_status()


def save_uploaded_logo(data: bytes, filename: str) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_LOGO_SUFFIXES:
        raise ValueError("Upload a PNG, JPG, WEBP, or GIF logo")
    if not data:
        raise ValueError("Empty upload")

    folder = images_dir()
    folder.mkdir(parents=True, exist_ok=True)

    # Clear previous logo.* variants
    for old in folder.glob("logo.*"):
        _archive_file(old)
        old.unlink()

    dest = folder / f"logo{suffix if suffix != '.jpeg' else '.jpg'}"
    if suffix == ".jpeg":
        dest = folder / "logo.jpg"
    dest.write_bytes(data)
    return template_status()


def apply_logo_to_gold_master(*, width_inches: float = 2.4) -> dict[str, Any]:
    """
    Insert the uploaded logo into the Word header of the gold master,
    then sync .docx + .dotx.
    """
    logo = logo_path()
    if logo is None:
        raise FileNotFoundError("No logo uploaded yet. Upload a logo first.")

    # Prefer editing the .docx twin, then convert to .dotx
    source = template_docx_path()
    if not source.exists():
        if template_dotx_path().exists():
            from .generate import _open_word_template

            doc = _open_word_template(template_dotx_path())
            doc.save(str(source))
        else:
            build_master_template(source)

    _archive_file(source)
    _archive_file(template_dotx_path())

    doc = Document(str(source))
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Wipe header content cleanly (old text + images)
    header_el = header._element
    for child in list(header_el):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            header_el.remove(child)

    p = header.add_paragraph()
    p.alignment = 0  # left
    run = p.add_run()
    run.add_picture(str(logo), width=Inches(width_inches))

    doc.save(str(source))
    docx_to_dotx(source, template_dotx_path())
    return template_status()


def rebuild_starter_with_optional_logo() -> dict[str, Any]:
    """Recreate the default gold master; apply logo if one is uploaded."""
    dest = template_docx_path()
    _archive_file(dest)
    _archive_file(template_dotx_path())
    build_master_template(dest)
    docx_to_dotx(dest, template_dotx_path())
    if logo_path() is not None:
        status = apply_logo_to_gold_master()
        _clear_body_masthead_when_logo_present()
        return template_status()
    return template_status()


def _clear_body_masthead_when_logo_present() -> None:
    """Avoid duplicating brand text in the body when the header logo already brands the page."""
    if logo_path() is None:
        return
    source = template_docx_path()
    if not source.exists():
        return
    doc = Document(str(source))
    changed = False
    if doc.paragraphs:
        first = doc.paragraphs[0].text.strip().upper()
        if first.startswith("SPOTLIGHT") or first.startswith("THE SPOTLIGHT"):
            for run in doc.paragraphs[0].runs:
                run.text = ""
            changed = True
        if len(doc.paragraphs) > 1 and "Official Document" in doc.paragraphs[1].text:
            for run in doc.paragraphs[1].runs:
                run.text = ""
            changed = True
    if changed:
        doc.save(str(source))
        docx_to_dotx(source, template_dotx_path())
