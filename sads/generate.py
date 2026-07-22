from __future__ import annotations

import io
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .brand import brand, brand_placeholders, primary_rgb, public_brand
from .clauses import ensure_default_clauses, expand_clause_references
from .components import (
    Component,
    assemble_body_components,
    resolve_image_path,
    table_rows_from_section,
)
from .models import Document, find_document_file, load_document
from .paths import (
    archive_dir,
    documents_root,
    load_config,
    output_docx_dir,
    output_pdf_dir,
    template_path,
)
from .template import DOCX_MAIN_CONTENT_TYPE, DOTX_MAIN_CONTENT_TYPE

NAVY = primary_rgb()
GRAY = RGBColor(0x55, 0x55, 0x55)
WARN = RGBColor(0x8B, 0x2E, 0x2E)
NOTE = RGBColor(0x3A, 0x4A, 0x5C)

LIBREOFFICE_CANDIDATES = (
    "soffice",
    "libreoffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/local/bin/soffice",
    "/opt/homebrew/bin/soffice",
)


def _open_word_template(path: Path) -> WordDocument:
    """
    Open a .docx or .dotx master template.

    python-docx only accepts document.main content types, so .dotx packages
    are remapped in memory before loading.
    """
    path = Path(path)
    if path.suffix.lower() != ".dotx":
        return WordDocument(str(path))

    buffer = io.BytesIO()
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(
        buffer, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "[Content_Types].xml":
                text = data.decode("utf-8").replace(
                    DOTX_MAIN_CONTENT_TYPE, DOCX_MAIN_CONTENT_TYPE
                )
                data = text.encode("utf-8")
            zout.writestr(info, data)
    buffer.seek(0)
    return WordDocument(buffer)


def _style_run(
    run,
    *,
    size: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    name: str = "Calibri",
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        text = paragraph.text
        for key, value in mapping.items():
            text = text.replace(key, value)
        if text != paragraph.text:
            paragraph.text = text
        return

    full = "".join(run.text for run in paragraph.runs)
    updated = full
    for key, value in mapping.items():
        updated = updated.replace(key, value)
    if updated == full:
        return

    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def _clear_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.text = ""


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _find_placeholder_paragraph(doc: WordDocument, placeholder: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def _write_paragraph(
    paragraph: Paragraph,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    space_after: int = 6,
) -> None:
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    _style_run(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)


def _heading_for_component(component: Component) -> tuple[str, RGBColor]:
    """Return display heading and color for a body component type."""
    stype = (component.type or "section").lower()
    label = {
        "warning": "Warning",
        "note": "Note",
        "tip": "Tip",
    }.get(stype)
    color = WARN if stype == "warning" else (NOTE if stype in ("note", "tip") else NAVY)
    if label and not component.heading.lower().startswith(label.lower()):
        return f"{label} — {component.heading}" if component.heading else label, color
    return component.heading, color


def _insert_table_after(doc: WordDocument, anchor: Paragraph, rows: list[list[str]]) -> Paragraph:
    """Insert a content table after anchor; return a spacer paragraph after the table."""
    if not rows:
        return anchor
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for r_i, row in enumerate(rows):
        for c_i in range(cols):
            value = row[c_i] if c_i < len(row) else ""
            _set_cell_text(table.rows[r_i].cells[c_i], value, bold=(r_i == 0), size=10)
    anchor._p.addnext(table._tbl)
    spacer_el = OxmlElement("w:p")
    table._tbl.addnext(spacer_el)
    return Paragraph(spacer_el, anchor._parent)


def _expand_body_placeholder(doc: WordDocument, document: Document) -> None:
    """Replace {{BODY}} with assembled typed building blocks."""
    target = _find_placeholder_paragraph(doc, "{{BODY}}")
    if target is None:
        return

    _clear_paragraph(target)
    components = assemble_body_components(
        purpose=document.purpose,
        scope=document.scope,
        sections=document.sections,
        include_purpose_scope=False,  # gold master already shows {{PURPOSE}}/{{SCOPE}}
    )
    if not components:
        return

    anchor = target
    first = True
    for component in components:
        stype = (component.type or "section").lower()

        # paragraph: body only (optional light heading)
        if stype == "paragraph":
            if component.heading:
                if first:
                    _write_paragraph(
                        anchor, component.heading, size=11, bold=True, color=NAVY, space_after=2
                    )
                    first = False
                else:
                    anchor = _insert_paragraph_after(anchor)
                    _write_paragraph(
                        anchor, component.heading, size=11, bold=True, color=NAVY, space_after=2
                    )
            body = expand_clause_references(component.body or "")
            for chunk in [c.strip() for c in body.split("\n\n") if c.strip()] or [""]:
                if first:
                    _write_paragraph(anchor, chunk, size=11, space_after=8)
                    first = False
                else:
                    anchor = _insert_paragraph_after(anchor)
                    _write_paragraph(anchor, chunk, size=11, space_after=8)
            continue

        heading, color = _heading_for_component(component)
        if heading:
            if first:
                _write_paragraph(anchor, heading, size=13, bold=True, color=color, space_after=4)
                first = False
            else:
                anchor = _insert_paragraph_after(anchor)
                _write_paragraph(anchor, heading, size=13, bold=True, color=color, space_after=4)

        if stype == "table":
            rows = table_rows_from_section(component)
            caption = expand_clause_references(component.body or "").strip()
            if rows and parse_looks_like_table(caption):
                caption = ""
            if caption:
                if first and not heading:
                    _write_paragraph(anchor, caption, size=10, color=GRAY, space_after=4)
                    first = False
                else:
                    anchor = _insert_paragraph_after(anchor)
                    _write_paragraph(anchor, caption, size=10, color=GRAY, space_after=4)
            if rows:
                if first:
                    first = False
                anchor = _insert_table_after(doc, anchor, rows)
            continue

        if stype == "image":
            caption = expand_clause_references(component.body or "").strip()
            image_path = resolve_image_path(component.src)
            if image_path is not None:
                if first and not heading:
                    run = anchor.add_run()
                    run.add_picture(str(image_path), width=Cm(14))
                    first = False
                else:
                    anchor = _insert_paragraph_after(anchor)
                    run = anchor.add_run()
                    run.add_picture(str(image_path), width=Cm(14))
            else:
                missing = f"[Image not found: {component.src or '(no src)'}]"
                if first and not heading:
                    _write_paragraph(anchor, missing, size=10, color=GRAY, space_after=6)
                    first = False
                else:
                    anchor = _insert_paragraph_after(anchor)
                    _write_paragraph(anchor, missing, size=10, color=GRAY, space_after=6)
            if caption:
                anchor = _insert_paragraph_after(anchor)
                _write_paragraph(anchor, caption, size=9, color=GRAY, space_after=8)
            continue

        body = expand_clause_references(component.body or "")
        body_chunks = [c.strip() for c in body.split("\n\n") if c.strip()]
        if not body_chunks:
            body_chunks = [body.strip()] if body.strip() else [""]

        for chunk in body_chunks:
            anchor = _insert_paragraph_after(anchor)
            if stype in ("signature", "approval"):
                _write_paragraph(anchor, chunk, size=10, space_after=10)
                anchor.paragraph_format.space_before = Pt(6)
            elif stype == "procedure":
                _write_paragraph(anchor, chunk, size=11, space_after=6)
            else:
                _write_paragraph(anchor, chunk, size=11, space_after=8)

        anchor = _insert_paragraph_after(anchor)
        _write_paragraph(anchor, "", size=11, space_after=4)


def parse_looks_like_table(text: str) -> bool:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    return len(lines) >= 2 and all("|" in ln for ln in lines[:3])


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _style_run(run, size=size, bold=bold, color=NAVY if bold else GRAY)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _expand_revision_placeholder(doc: WordDocument, document: Document) -> None:
    """Replace {{REVISION_HISTORY}} with a structured Version/Date/Author/Description table."""
    target = _find_placeholder_paragraph(doc, "{{REVISION_HISTORY}}")
    if target is None:
        return

    _clear_paragraph(target)
    revisions = document.revision_history
    if not revisions:
        _write_paragraph(
            target,
            "No revisions recorded.",
            size=9,
            color=GRAY,
            space_after=4,
        )
        return

    # Build table at end of document, then move after the placeholder paragraph
    table = doc.add_table(rows=1 + len(revisions), cols=4)
    table.style = "Table Grid"
    table.autofit = True

    headers = ("Version", "Date", "Author", "Description")
    for i, label in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], label, bold=True, size=9)

    for row_i, rev in enumerate(revisions, start=1):
        _set_cell_text(table.rows[row_i].cells[0], rev.version)
        _set_cell_text(table.rows[row_i].cells[1], rev.date)
        _set_cell_text(table.rows[row_i].cells[2], rev.author)
        _set_cell_text(table.rows[row_i].cells[3], rev.notes)

    # Prefer reasonable column widths (deterministic layout)
    widths = (Cm(2.2), Cm(2.6), Cm(3.4), Cm(8.0))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    target._p.addnext(table._tbl)

def _replace_placeholders(doc: WordDocument, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                _replace_in_paragraph(paragraph, mapping)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                _replace_in_paragraph(paragraph, mapping)


def _format_document_info(document: Document) -> str:
    return (
        f"Version {document.version}  ·  Category: {document.category}  ·  "
        f"Owner: {document.owner}  ·  Status: {document.approved}"
    )


def _archive_existing(path: Path) -> None:
    if not path.exists():
        return
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = archive_dir() / f"{path.stem}-{stamp}{path.suffix}"
    archive_dir().mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, target)


def pdf_backends_available() -> list[str]:
    """Return human-readable names of PDF converters found on this machine."""
    found: list[str] = []
    # Presence of Word.app does not guarantee an activated/licensed install.
    word = Path("/Applications/Microsoft Word.app")
    if word.exists():
        found.append("Microsoft Word.app (may require license)")
    for binary in LIBREOFFICE_CANDIDATES:
        path = Path(binary) if binary.startswith("/") else shutil.which(binary)
        if path and Path(path).exists():
            found.append(f"LibreOffice ({path})")
            break
    return found


def _export_pdf(docx_path: Path, pdf_path: Path) -> tuple[Path | None, str | None]:
    """
    Export DOCX → PDF.

    Returns (pdf_path, None) on success, or (None, reason) when skipped/unavailable.
    """
    cfg = load_config().get("pdf", {})
    if not cfg.get("enabled", True):
        return None, "PDF disabled in config.json"

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    method = cfg.get("method", "auto")

    # Microsoft Word via AppleScript (works with full Word installs)
    if method in ("auto", "word"):
        word_script = f'''
        tell application "Microsoft Word"
          set theDoc to open POSIX file "{docx_path}"
          save as theDoc file name POSIX file "{pdf_path}" file format format PDF
          close theDoc saving no
        end tell
        '''
        try:
            result = subprocess.run(
                ["osascript", "-e", word_script],
                capture_output=True,
                text=True,
                check=False,
                timeout=120,
            )
            if result.returncode == 0 and pdf_path.exists():
                return pdf_path, None
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

    # LibreOffice / soffice (free; preferred when Word is unpaid / unavailable)
    if method in ("auto", "libreoffice"):
        for binary in LIBREOFFICE_CANDIDATES:
            resolved = binary if binary.startswith("/") else shutil.which(binary)
            if not resolved or not Path(resolved).exists():
                continue
            try:
                result = subprocess.run(
                    [
                        str(resolved),
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(pdf_path.parent),
                        str(docx_path),
                    ],
                    capture_output=True,
                    text=True,
                    check=False,
                    timeout=180,
                )
                produced = pdf_path.parent / f"{docx_path.stem}.pdf"
                if result.returncode == 0 and produced.exists():
                    if produced != pdf_path:
                        produced.replace(pdf_path)
                    return pdf_path, None
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

    return (
        None,
        "No PDF converter found. Install free LibreOffice, or use --no-pdf. "
        "Then: brew install --cask libreoffice",
    )


def generate(number: str, *, make_pdf: bool = True) -> dict:
    """
    Phase 4 pipeline:

    Open template → replace placeholders → expand body/revisions →
    save DOCX → optional PDF → done.
    """
    number = number.strip().upper()
    ensure_default_clauses()
    source = find_document_file(number, documents_root())
    document = load_document(source)
    template = template_path()

    mapping = {
        "{{NUMBER}}": document.number,
        "{{TITLE}}": document.title,
        "{{PURPOSE}}": expand_clause_references(document.purpose),
        "{{SCOPE}}": expand_clause_references(document.scope),
        "{{VERSION}}": document.version,
        "{{CATEGORY}}": document.category,
        "{{OWNER}}": document.owner,
        "{{APPROVED}}": document.approved,
        "{{DOCUMENT_INFO}}": _format_document_info(document),
        # Expanded separately; leave empty if a stray token remains
        "{{REVISION_HISTORY}}": "",
        "{{BODY}}": "",
        **brand_placeholders(),
        # Legacy legal entity rename (global)
        "Spotlight Alliance LLC": brand()["legal_company"],
        "Spotlight Alliance": brand()["legal_company"],
    }

    out_docx = output_docx_dir() / f"{document.number}.docx"
    out_pdf = output_pdf_dir() / f"{document.number}.pdf"
    output_docx_dir().mkdir(parents=True, exist_ok=True)

    _archive_existing(out_docx)
    if make_pdf:
        _archive_existing(out_pdf)

    word = _open_word_template(template)
    _expand_body_placeholder(word, document)
    _expand_revision_placeholder(word, document)
    _replace_placeholders(word, mapping)
    word.save(str(out_docx))

    pdf_result: Path | None = None
    pdf_note: str | None = None
    if make_pdf:
        pdf_result, pdf_note = _export_pdf(out_docx, out_pdf)
    else:
        pdf_note = "skipped (--no-pdf)"

    return {
        "number": document.number,
        "docx": out_docx,
        "pdf": pdf_result,
        "pdf_note": pdf_note,
        "source": source,
        "template": template,
    }
