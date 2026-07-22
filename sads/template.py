from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .brand import brand, primary_rgb, public_brand, tagline


NAVY = RGBColor(0x1A, 0x2B, 0x4A)
GRAY = RGBColor(0x55, 0x55, 0x55)
RULE = "1A2B4A"

# OOXML: .docx vs .dotx differs mainly by the main document content type.
DOCX_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.document.main+xml"
)
DOTX_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.template.main+xml"
)

# Placeholders that must appear in the gold master (Phase 2 contract).
REQUIRED_PLACEHOLDERS = (
    "{{NUMBER}}",
    "{{TITLE}}",
    "{{DOCUMENT_INFO}}",
    "{{PURPOSE}}",
    "{{SCOPE}}",
    "{{BODY}}",
    "{{REVISION_HISTORY}}",
    "{{VERSION}}",
)


def docx_to_dotx(source: Path, destination: Path | None = None) -> Path:
    """
    Convert a .docx package to a .dotx template without Microsoft Word.

    A .dotx is the same ZIP/Open XML package as a .docx, with the main
    document content type set to template.main+xml instead of document.main+xml.
    """
    source = Path(source)
    if not source.exists():
        raise FileNotFoundError(f"Source document not found: {source}")
    if destination is None:
        destination = source.with_suffix(".dotx")
    else:
        destination = Path(destination)

    destination.parent.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(source, "r") as zin:
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    if DOTX_MAIN_CONTENT_TYPE not in text:
                        if DOCX_MAIN_CONTENT_TYPE not in text:
                            raise ValueError(
                                "Source is not a WordprocessingML document "
                                f"(missing document.main content type): {source}"
                            )
                        text = text.replace(
                            DOCX_MAIN_CONTENT_TYPE, DOTX_MAIN_CONTENT_TYPE
                        )
                    data = text.encode("utf-8")
                elif info.filename == "word/settings.xml":
                    data = _ensure_template_document_type(data)
                elif info.filename == "docProps/app.xml":
                    data = _set_template_name(data, destination.name)
                zout.writestr(info, data)

    destination.write_bytes(buffer.getvalue())
    return destination


def _ensure_template_document_type(settings_xml: bytes) -> bytes:
    """Mark the package as a template in word/settings.xml when possible."""
    text = settings_xml.decode("utf-8")
    if "w:documentType" in text:
        text = re.sub(
            r"<w:documentType\s+[^/]*?/>",
            '<w:documentType w:val="template"/>',
            text,
            count=1,
        )
        return text.encode("utf-8")

    match = re.search(r"<w:settings\b[^>]*>", text)
    if not match:
        return settings_xml
    insert_at = match.end()
    text = (
        text[:insert_at]
        + '<w:documentType w:val="template"/>'
        + text[insert_at:]
    )
    return text.encode("utf-8")


def _set_template_name(app_xml: bytes, template_filename: str) -> bytes:
    text = app_xml.decode("utf-8")
    if re.search(r"<Template>.*?</Template>", text, flags=re.DOTALL):
        text = re.sub(
            r"<Template>.*?</Template>",
            f"<Template>{template_filename}</Template>",
            text,
            count=1,
            flags=re.DOTALL,
        )
    else:
        text = text.replace(
            "</Properties>",
            f"<Template>{template_filename}</Template></Properties>",
            1,
        )
    return text.encode("utf-8")


def _set_run(
    run,
    *,
    size: int,
    bold: bool = False,
    color: RGBColor | None = None,
    name: str = "Calibri",
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color


def _space_after(paragraph, points: float) -> None:
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.space_before = Pt(0)


def _bottom_rule(paragraph, *, color: str = RULE, size: str = "12") -> None:
    """Add a bottom border under a paragraph (brand rule)."""
    p_pr = paragraph._p.get_or_add_pPr()
    # Remove existing borders if re-running builder logic
    for child in list(p_pr):
        if child.tag == qn("w:pBdr"):
            p_pr.remove(child)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE field into an existing paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def build_master_template(destination: Path) -> Path:
    """
    Create the Phase 2 gold-master Word template with locked placeholders.

    Formatting lives here once. Content never lives here — only {{PLACEHOLDERS}}.
    """
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    identity = brand()
    navy = primary_rgb()
    rule_hex = identity["primary_hex"].lstrip("#").upper()

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)

    # --- Brand masthead ---
    brand_line = doc.add_paragraph()
    _space_after(brand_line, 2)
    run = brand_line.add_run(public_brand().upper())
    _set_run(run, size=12, bold=True, color=navy)

    tagline_p = doc.add_paragraph()
    _space_after(tagline_p, 6)
    run = tagline_p.add_run(tagline())
    _set_run(run, size=9, color=GRAY)
    _bottom_rule(tagline_p, color=rule_hex)

    spacer = doc.add_paragraph()
    _space_after(spacer, 10)

    # --- Identity ---
    number = doc.add_paragraph()
    _space_after(number, 2)
    run = number.add_run("{{NUMBER}}")
    _set_run(run, size=11, bold=True, color=navy)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = title.add_run("{{TITLE}}")
    _set_run(run, size=22, bold=True, color=navy)

    meta = doc.add_paragraph()
    _space_after(meta, 14)
    run = meta.add_run("{{DOCUMENT_INFO}}")
    _set_run(run, size=9, color=GRAY)

    # --- Purpose ---
    purpose_h = doc.add_paragraph()
    _space_after(purpose_h, 4)
    run = purpose_h.add_run("Purpose")
    _set_run(run, size=13, bold=True, color=navy)

    purpose = doc.add_paragraph()
    _space_after(purpose, 12)
    purpose.paragraph_format.line_spacing = 1.15
    run = purpose.add_run("{{PURPOSE}}")
    _set_run(run, size=11)

    # --- Scope ---
    scope_h = doc.add_paragraph()
    _space_after(scope_h, 4)
    run = scope_h.add_run("Scope")
    _set_run(run, size=13, bold=True, color=navy)

    scope = doc.add_paragraph()
    _space_after(scope, 12)
    scope.paragraph_format.line_spacing = 1.15
    run = scope.add_run("{{SCOPE}}")
    _set_run(run, size=11)

    # --- Body (sections expanded by generator) ---
    body = doc.add_paragraph()
    _space_after(body, 12)
    body.paragraph_format.line_spacing = 1.15
    run = body.add_run("{{BODY}}")
    _set_run(run, size=11)

    # --- Revision History ---
    rev_h = doc.add_paragraph()
    _space_after(rev_h, 4)
    run = rev_h.add_run("Revision History")
    _set_run(run, size=13, bold=True, color=navy)

    rev = doc.add_paragraph()
    _space_after(rev, 6)
    run = rev.add_run("{{REVISION_HISTORY}}")
    _set_run(run, size=9, color=GRAY)

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear default empty paragraph content
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(fp.runs):
        run.text = ""
    left = fp.add_run(
        f"© {identity['copyright']}  ·  {identity['public_brand']}  ·  "
        f"{identity['website']}  ·  Confidential  ·  {{{{NUMBER}}}}  ·  "
        f"Rev {{{{VERSION}}}}  ·  Page "
    )
    _set_run(left, size=8, color=GRAY)
    _add_page_number_field(fp)
    for run in fp.runs:
        if run.font.size is None:
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY
            run.font.name = "Calibri"

    # Tighten default style a touch for consistency
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.save(destination)
    return destination


def build_master_template_pair(docx_path: Path) -> tuple[Path, Path]:
    """Create the .docx working copy and matching .dotx gold master."""
    docx_path = build_master_template(docx_path)
    dotx_path = docx_to_dotx(docx_path, docx_path.with_suffix(".dotx"))
    return docx_path, dotx_path


def template_text_blob(path: Path) -> str:
    """Collect visible text from a template for placeholder validation."""
    from .generate import _open_word_template

    doc = _open_word_template(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    parts.append(paragraph.text)
    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            for paragraph in header.paragraphs:
                parts.append(paragraph.text)
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in footer.paragraphs:
                parts.append(paragraph.text)
    return "\n".join(parts)


def validate_template(path: Path | None = None) -> dict:
    """
    Confirm the gold master exists and contains every required placeholder.
    """
    from .paths import template_path

    target = Path(path) if path is not None else template_path()
    if not target.exists():
        raise FileNotFoundError(f"Gold master not found: {target}")

    blob = template_text_blob(target)
    missing = [ph for ph in REQUIRED_PLACEHOLDERS if ph not in blob]
    return {
        "path": target,
        "ok": not missing,
        "missing": missing,
        "is_dotx": target.suffix.lower() == ".dotx",
    }
